from copy import deepcopy
from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


SOURCE = Path(r"C:\Users\pc\Desktop\办公\打分标准.docx")
OUTPUT = Path(r"C:\Users\pc\Desktop\办公\打分标准_优化版.docx")


def set_run_font(run, name, size=None, bold=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def set_style_font(style, name, size, bold=False, color=(0, 0, 0)):
    style.font.name = name
    style._element.rPr.rFonts.set(qn("w:ascii"), name)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor(*color)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, "Microsoft YaHei", 10.5)
    normal.paragraph_format.line_spacing = 1.45
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.widow_control = True

    title = styles["Title"]
    set_style_font(title, "Microsoft YaHei", 22, True, (31, 78, 121))
    title.paragraph_format.space_after = Pt(12)
    title.paragraph_format.line_spacing = 1.1

    heading_tokens = {
        "Heading 1": (16, (31, 78, 121), 18, 8),
        "Heading 2": (13, (46, 92, 133), 12, 5),
        "Heading 3": (11.5, (47, 47, 47), 9, 3),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = styles[style_name]
        set_style_font(style, "Microsoft YaHei", size, True, color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
        style.paragraph_format.line_spacing = 1.2

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        set_style_font(style, "Microsoft YaHei", 10.5)
        style.paragraph_format.line_spacing = 1.4
        style.paragraph_format.space_after = Pt(3)

    if "Note" not in styles:
        note = styles.add_style("Note", WD_STYLE_TYPE.PARAGRAPH)
    else:
        note = styles["Note"]
    set_style_font(note, "Microsoft YaHei", 10.5, False, (71, 71, 71))
    note.paragraph_format.left_indent = Cm(0.55)
    note.paragraph_format.first_line_indent = Cm(0)
    note.paragraph_format.space_after = Pt(6)
    note.paragraph_format.line_spacing = 1.4


def remove_all_body(doc):
    body = doc._element.body
    for element in list(body):
        if element.tag != qn("w:sectPr"):
            body.remove(element)


def add_markdown_runs(paragraph, text, size=None, color=None):
    for fragment in re.split(r"(\*\*.*?\*\*)", text):
        if not fragment:
            continue
        is_bold = fragment.startswith("**") and fragment.endswith("**")
        run = paragraph.add_run(fragment[2:-2] if is_bold else fragment)
        set_run_font(run, "Microsoft YaHei", size=size, bold=is_bold, color=color)


def add_footer(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("打分标准  |  第 ")
    set_run_font(run, "Microsoft YaHei", 8.5, color=(112, 112, 112))
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    cached = OxmlElement("w:t")
    cached.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, cached, end])
    tail = paragraph.add_run(" 页")
    set_run_font(tail, "Microsoft YaHei", 8.5, color=(112, 112, 112))


def is_subheading(lines, index, content):
    if not re.match(r"^\d+\.\s+", content):
        return False
    for following in lines[index + 1:]:
        candidate = following.strip()
        if not candidate:
            continue
        return candidate.startswith("-") or candidate.startswith("   ")
    return False


def main():
    source = Document(SOURCE)
    lines = [p.text for p in source.paragraphs]
    doc = Document()
    section = doc.sections[0]
    # Preserve the source document's A4 page form while giving the text a
    # slightly wider reading column than the original's very large margins.
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.35)
    section.right_margin = Cm(2.35)
    configure_styles(doc)
    remove_all_body(doc)
    add_footer(section)

    for index, raw in enumerate(lines):
        text = raw.strip()
        if not text:
            continue

        if text.startswith("# "):
            content = text[2:].strip()
            style = "Title" if index == 0 else "Heading 1"
            paragraph = doc.add_paragraph(style=style)
            add_markdown_runs(paragraph, content, size=22 if style == "Title" else 16,
                              color=(31, 78, 121))
            if style == "Title":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        if text.startswith("## "):
            paragraph = doc.add_paragraph(style="Heading 1")
            add_markdown_runs(paragraph, text[3:].strip(), size=16, color=(31, 78, 121))
            continue

        if text.startswith("### "):
            paragraph = doc.add_paragraph(style="Heading 2")
            add_markdown_runs(paragraph, text[4:].strip(), size=13, color=(46, 92, 133))
            continue

        if is_subheading(lines, index, text):
            paragraph = doc.add_paragraph(style="Heading 3")
            add_markdown_runs(paragraph, text, size=11.5, color=(47, 47, 47))
            continue

        if text.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            add_markdown_runs(paragraph, text[2:].strip(), size=10.5)
            continue

        numbered = re.match(r"^(\d+)\.\s+(.+)$", text)
        if numbered:
            paragraph = doc.add_paragraph(style="List Number")
            add_markdown_runs(paragraph, numbered.group(2), size=10.5)
            continue

        if raw.startswith("   "):
            paragraph = doc.add_paragraph(style="Note")
            add_markdown_runs(paragraph, text, size=10.5)
            continue

        paragraph = doc.add_paragraph(style="Normal")
        add_markdown_runs(paragraph, text, size=10.5)

    doc.core_properties.title = "通用学习/工作文档完整百分制打分标准"
    doc.core_properties.subject = "格式优化版"
    doc.core_properties.author = ""
    doc.save(OUTPUT)


if __name__ == "__main__":
    main()
